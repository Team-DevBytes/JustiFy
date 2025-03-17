from flask import Flask, request, jsonify, render_template, session, send_file
from openai import OpenAI
import PyPDF2
import os 
from dotenv import load_dotenv
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import datetime
from multiagent import *

load_dotenv()

api_key = os.getenv('OPENAI_API_KEY')

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'default-secret-key')

general_context = ""
doc_chat_context = ""

# OpenAI API Client Configuration
client = OpenAI(api_key=api_key)

# Store document text globally (in a real app, you'd use a database or session)
document_cache = {}
pdf_cache = {}  # Store PDF files for viewing
draft_cache = {}  # Store generated drafts

CATEGORY_METRICS = {
    'Legal Notice': [
        'Severity Score', 'Violations & Broken Rules', 'Legal Consequences', 'Actionable Steps',
        'Urgency Detection', 'Tone Analysis', 'Recommended Actions'
    ],
    'Ownership Documents': [
        'Ownership Rights & Obligations', 'Transfer, Leasing, Sale, Mortgaging Clauses',
        'Financial Liabilities', 'Terms & Conditions', 'Important Dates', 'Document Validity', 'Summary Type'
    ],
    'Contracts & Agreements': [
        'Parties Involved & Roles', 'Terms & Conditions', 'Termination Clauses', 'Penalties for Breach',
        'Severity Score', 'Obligations & Rights', 'Actionable Steps'
    ],
    'Financial Documents': [
        'Financial Obligations', 'Coverage Details', 'Deadlines & Payment Schedules',
        'Legal Implications', 'Severity Score', 'Urgency Detection', 'Risk Analysis'
    ],
    'Terms & Conditions / Privacy Policies': [
        'User Rights & Restrictions', 'Data Usage & Privacy Clauses', 'Liability Clauses',
        'Termination & Suspension Rules', 'Severity Score', 'Personal Implications', 'Suggested Actions'
    ],
    'Intellectual Property Documents': [
        'Ownership & Usage Rights', 'Infringement Clauses', 'Exclusivity & Licensing Terms',
        'Penalties for Violation', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Criminal Offense Documents': [
        'Charges Filed', 'Potential Penalties', 'Required Actions', 'Severity Score',
        'Urgency Detection', 'Tone Analysis', 'Suggested Actions'
    ],
    'Regulatory Compliance Documents': [
        'Compliance Requirements', 'Penalties for Non-Compliance', 'Renewal Deadlines & Conditions',
        'Guidelines for Rectification', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Employment Documents': [
        'Terms of Employment', 'Termination Conditions', 'Confidentiality Clauses',
        'Breach Consequences', 'Severity Score', 'Urgency Detection', 'Suggested Actions'
    ],
    'Court Judgments & Legal Precedents': [
        'Summary of Judgment', 'Legal Basis', 'Potential Consequences',
        'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ]
}

# Document templates for different draft types
DRAFT_TEMPLATES = {
    'Legal Notice Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Contract Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Arial', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'date_format': '%d/%m/%Y',  # Example: 01/01/2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'General Letter': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Calibri', 'size': 12, 'bold': True, 'align': 'left'},
        'body_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Legal Memo': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',  # Example: January 1, 2023
        'includes_header': True,
        'includes_date': True,
        'includes_signature': False
    }
}

# Map document categories to draft templates
CATEGORY_TO_TEMPLATE = {
    'Legal Notice': 'Legal Notice Response',
    'Contracts & Agreements': 'Contract Response',
    'Ownership Documents': 'Legal Memo',
    'Financial Documents': 'Legal Memo',
    'Terms & Conditions / Privacy Policies': 'Legal Memo',
    'Intellectual Property Documents': 'Legal Memo',
    'Criminal Offense Documents': 'Legal Notice Response',
    'Regulatory Compliance Documents': 'Legal Memo',
    'Employment Documents': 'Legal Memo',
    'Court Judgments & Legal Precedents': 'Legal Memo',
    'default': 'General Letter'
}

@app.route('/')
def index():
    # Generate a unique session ID if not exists
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()
    return render_template('index.html')

@app.route('/general_chat.html')
def general_chat():
    return render_template('general_chat.html')

def extract_text_from_pdf(pdf_file):
    try:
        # Create a BytesIO object to avoid file seeking issues
        pdf_content = BytesIO(pdf_file.read())
        # Reset the file pointer for future operations
        pdf_file.seek(0)
        
        reader = PyPDF2.PdfReader(pdf_content)
        text = ''
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
        return text
    except Exception as e:
        app.logger.error(f"PDF extraction error: {str(e)}")
        return str(e)


@app.route('/classify', methods=['POST'])
def classify_document():
    if 'document' not in request.files:
        return jsonify({'error': 'No PDF file uploaded'}), 400

    pdf_file = request.files['document']
    
    # Save the PDF file in memory for later viewing
    session_id = session.get('session_id', os.urandom(16).hex())
    pdf_content = pdf_file.read()
    pdf_cache[session_id] = pdf_content
    
    # Reset file pointer
    pdf_file.seek(0)
    
    document_text = extract_text_from_pdf(pdf_file)

    if not document_text:
        return jsonify({'error': 'Failed to extract text from PDF'}), 400

    # Save the document text in the cache using session ID
    document_cache[session_id] = document_text

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a document classification agent. Classify the document into one of these categories: Legal Notice, Ownership Documents, Contracts & Agreements, Financial Documents, Terms & Conditions / Privacy Policies, Intellectual Property Documents, Criminal Offense Documents, Regulatory Compliance Documents, Employment Documents, Court Judgments & Legal Precedents."},
                {"role": "user", "content": document_text[:3000]}  # Limit to first 3000 characters
            ]
        )

        category = response.choices[0].message.content.strip()
        return jsonify({'category': category})
    except Exception as e:
        app.logger.error(f"Classification error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/process', methods=['POST'])
def process_document():
    if 'document' not in request.files or 'category' not in request.form:
        return jsonify({'error': 'Document file or category is missing'}), 400

    pdf_file = request.files['document']
    category = request.form['category']
    document_text = extract_text_from_pdf(pdf_file)

    if not document_text or not category:
        return jsonify({'error': 'Document text or category is missing'}), 400

    # Save the document text in the cache using session ID
    session_id = session.get('session_id', os.urandom(16).hex())
    document_cache[session_id] = document_text

    metrics = CATEGORY_METRICS.get(category, [])
    metrics_prompt = ', '.join(metrics)

    prompt = f"You are an expert summarizer for {category} documents. Extract the following relevant metrics: {metrics_prompt}. Format each metric as '**Metric Name**: Value' to make it bold and easily readable."

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": document_text[:4000]}  # Limit text to avoid token limits
            ]
        )

        summary = response.choices[0].message.content.strip()
        # Return both the summary and the document text (truncated for frontend)
        return jsonify({
            'summary': summary,
            'document_text': document_text[:200] + '...' if len(document_text) > 200 else document_text
        })
    except Exception as e:
        app.logger.error(f"Processing error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/chat', methods=['POST'])
def chat():
    global doc_chat_context
    data = request.json
    user_message = data.get('message')
    category = data.get('category')
    detailed_analysis = data.get('detailed_analysis', False)
    generate_draft = data.get('generate_draft', False)
    draft_instructions = data.get('draft_instructions', '')
    
    if not user_message:
        return jsonify({'error': 'Message is required'}), 400
    
    # Get the document text from the cache
    session_id = session.get('session_id')
    document_text = document_cache.get(session_id, '')
    
    if not document_text:
        return jsonify({'error': 'No document found. Please process a document first.'}), 400
    
    try:
        # Prepare document context
        document_context = document_text[:3000] if len(document_text) > 3000 else document_text
        
        # Construct a system prompt that includes document information and category
        system_prompt = f"""You are a legal assistant specializing in {category} documents.
        You will answer only the questions related to the document and not any external questions like generating code or writing a story.
        You have access to the following document text (truncated if necessary):

        {document_context}

        chat Context:
        {doc_chat_context}

"""
        
        # Add instructions based on detailed analysis or draft generation
        if detailed_analysis:
            system_prompt += "Provide a detailed analysis with comprehensive explanations, legal references, and thorough examination of all relevant aspects. "
        else:
            system_prompt += "Provide concise, clear answers focused on the most important points. "
            
        if generate_draft:
            # For draft generation, we'll handle it separately
            draft_id = generate_document_draft(user_message, draft_instructions, category, document_context)
            return jsonify({
                'response': f"I've prepared a draft document based on your instructions. You can download it using the link below.",
                'draft_id': draft_id
            })
        
        system_prompt += "Provide helpful, accurate information based on this document. If you cannot find information in the document to answer a question, clearly state that. Use **bold** for important points."
        
        # Make API call to OpenAI
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ]
        )
        bot_response = response.choices[0].message.content.strip()
        doc_chat_context += f"\nUser: {user_message}\n"
        doc_chat_context += f"\nBot: {bot_response}\n"
        return jsonify({'response': bot_response})
    
    except Exception as e:
        app.logger.error(f"Chat error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/general_chat', methods=['POST'])
def general_chat_api():
    global general_context
    data = request.json
    user_message = data.get('message')
    detailed_analysis = data.get('detailed_analysis', False)
    
    if not user_message:
        return jsonify({'error': 'Message is required'}), 400
    
    try:
        if detailed_analysis:
            response, reasoning = get_answer(user_message, general_context)
            general_context += f"\n User: {user_message}\n"
            general_context += f"\nSenior Lawyer: {response}\n"
            return jsonify({'response': response, 'reasoning': reasoning})

        else:
            system_prompt = f"""You are a knowledgeable legal assistant who can provide general information about legal topics. 
            You are not a lawyer and should clarify that your responses do not constitute legal advice. 
            You should recommend consulting with a qualified attorney for specific lexgal situations.
            Provide concise, clear answers focused on the most important points.
            Use **bold** for important points and structure your response in a clear, organized manner.
            
            Context: {general_context}
            """     
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ]
            )
            bot_response = response.choices[0].message.content.strip()
            return jsonify({'response': bot_response, 'reasoning': []})
        
    except Exception as e:
        app.logger.error(f"General chat error: {str(e)}")
        return jsonify({'error': str(e)}), 500


def generate_document_draft(message, instructions, category, document_context):
    """Generate a formatted document draft based on the document category and user instructions"""
    
    # Determine the template to use based on the document category
    template_name = CATEGORY_TO_TEMPLATE.get(category, CATEGORY_TO_TEMPLATE['default'])
    template = DRAFT_TEMPLATES[template_name]
    
    # Create a prompt for the draft generation
    prompt = f"""You are a professional legal document drafter. Create a formal response document based on the following instructions:

Instructions: {instructions}

This is related to a {category} document. Here's the relevant context from the document:
{document_context[:1500]}

Your draft should be well-structured and professionally formatted. Include:
1. A clear header/title
2. Today's date ({datetime.datetime.now().strftime(template['date_format'])})
3. Appropriate salutation
4. Well-organized body content
5. Proper closing
6. Signature line

Format the content as if it were going to be printed on letterhead. Do not include any explanatory text or notes - just the actual document content.
"""

    try:
        # Generate the draft content using OpenAI
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": message}
            ]
        )
        
        draft_content = response.choices[0].message.content.strip()
        
        # Create a formatted Word document
        doc = create_formatted_document(draft_content, template)
        
        # Save the document to a temporary file
        draft_id = str(uuid.uuid4())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # Store the file path in the draft cache
        draft_cache[draft_id] = {
            'path': temp_file.name,
            'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        }
        
        return draft_id
        
    except Exception as e:
        app.logger.error(f"Draft generation error: {str(e)}")
        raise


def generate_general_draft(message, instructions):
    """Generate a formatted document draft for general legal inquiries"""
    
    # Use the general letter template for general drafts
    template = DRAFT_TEMPLATES['General Letter']
    
    # Create a prompt for the draft generation
    prompt = f"""You are a professional legal document drafter. Create a formal document based on the following instructions:

Instructions: {instructions}

Your draft should be well-structured and professionally formatted. Include:
1. A clear header/title
2. Today's date ({datetime.datetime.now().strftime(template['date_format'])})
3. Appropriate salutation (if applicable)
4. Well-organized body content
5. Proper closing
6. Signature line (if applicable)

Format the content as if it were going to be printed on letterhead. Do not include any explanatory text or notes - just the actual document content.
"""

    try:
        # Generate the draft content using OpenAI
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": message}
            ]
        )
        
        draft_content = response.choices[0].message.content.strip()
        
        # Create a formatted Word document
        doc = create_formatted_document(draft_content, template)
        
        # Save the document to a temporary file
        draft_id = str(uuid.uuid4())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # Store the file path in the draft cache
        draft_cache[draft_id] = {
            'path': temp_file.name,
            'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        }
        
        return draft_id
        
    except Exception as e:
        app.logger.error(f"Draft generation error: {str(e)}")
        raise


def create_formatted_document(content, template):
    """Create a properly formatted Word document based on the template and content"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(template['margins']['top'])
        section.bottom_margin = Inches(template['margins']['bottom'])
        section.left_margin = Inches(template['margins']['left'])
        section.right_margin = Inches(template['margins']['right'])
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line and apply appropriate formatting
    current_section = 'header'
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
        
        # Determine the section based on content
        if current_section == 'header' and template['includes_date'] and any(month in line.lower() for month in ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']) and any(str(day) in line for day in range(1, 32)):
            current_section = 'date'
        elif current_section in ['header', 'date'] and any(salutation in line.lower() for salutation in ['dear', 'to whom', 'attention', 'attn', 're:', 'subject:']):
            current_section = 'salutation'
        elif current_section in ['header', 'date', 'salutation'] and any(closing in line.lower() for closing in ['sincerely', 'regards', 'truly', 'thank you', 'best', 'respectfully']):
            current_section = 'signature'
        elif current_section in ['header', 'date', 'salutation'] and len(line) > 20:
            current_section = 'body'
        
        # Apply formatting based on the section
        p = doc.add_paragraph()
        
        if current_section == 'header':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template['header_format']['align'] == 'center' else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['header_format']['font']
            run.font.size = Pt(template['header_format']['size'])
            run.font.bold = template['header_format']['bold']
        
        elif current_section == 'date':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['body_format']['font']
            run.font.size = Pt(template['body_format']['size'])
        
        elif current_section == 'body':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['body_format']['font']
            run.font.size = Pt(template['body_format']['size'])
        
        elif current_section == 'signature':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['signature_format']['font']
            run.font.size = Pt(template['signature_format']['size'])
    
    return doc


@app.route('/download-draft/<draft_id>', methods=['GET'])
def download_draft(draft_id):
    """Download a generated draft document"""
    
    if draft_id not in draft_cache:
        return jsonify({'error': 'Draft not found'}), 404
    
    draft_info = draft_cache[draft_id]
    
    try:
        return send_file(
            draft_info['path'],
            as_attachment=True,
            download_name=draft_info['filename'],
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        app.logger.error(f"Draft download error: {str(e)}")
        return jsonify({'error': 'Error downloading draft'}), 500


@app.route('/view-document', methods=['GET'])
def view_document():
    session_id = session.get('session_id')
    if not session_id or session_id not in pdf_cache:
        return jsonify({'error': 'No document found'}), 404
    
    # Create a temporary file to serve
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.write(pdf_cache[session_id])
    temp_file.close()
    
    return send_file(temp_file.name, mimetype='application/pdf', as_attachment=False)


if __name__ == '__main__':
    app.run(debug=True)

